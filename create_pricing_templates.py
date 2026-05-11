"""
Creates 3 pricing proposal Word templates for the ERP contracts module.
Run once, then upload each file via Contracts → Manage Templates.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT = os.path.dirname(os.path.abspath(os.getcwd() + "/create_pricing_templates.py"))

NAVY   = RGBColor(0x0A, 0x2A, 0x52)
GREEN  = RGBColor(0x3A, 0x6B, 0x35)
BLUE   = RGBColor(0x1A, 0x5F, 0x9E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF4, 0xF8, 0xFF)
MUTED  = RGBColor(0x55, 0x65, 0x7A)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D0D8E8"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),  "single")
        el.set(qn("w:sz"),   "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def para(cell_or_doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    p = cell_or_doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def header_cell(cell, text, bg="0A2A52"):
    set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE

def value_cell(cell, label, value, bg="F4F8FF"):
    set_cell_bg(cell, bg)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    lr = p.add_run(label + "\n")
    lr.font.size = Pt(9)
    lr.font.color.rgb = MUTED
    vr = p.add_run(value)
    vr.bold = True
    vr.font.size = Pt(13)
    vr.font.color.rgb = NAVY

def make_doc():
    doc = Document()
    # Narrow margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
    # Remove default styles noise
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    return doc

def add_title(doc, title="PRICING"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

def add_client(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run("@@party2_name@@")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = GREEN

def add_notes_footer(doc):
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "EFF4FB")
    set_cell_borders(cell, "C5D3E8")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    lr = p.add_run("Notes:  ")
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = NAVY
    nr = p.add_run("@@notes@@")
    nr.font.size = Pt(10)
    nr.font.color.rgb = MUTED

    doc.add_paragraph()
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = pd.add_run("Date: @@contract_date@@")
    dr.font.size = Pt(10)
    dr.font.color.rgb = MUTED

def add_bonus_note(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("* Success premiums will be calculated based on the benefit expected following project approval.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


# ─── TEMPLATE 1 — Single program ──────────────────────────────────
doc1 = make_doc()
add_title(doc1)
add_client(doc1)

tbl1 = doc1.add_table(rows=3, cols=1)
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl1.style = "Table Grid"
header_cell(tbl1.cell(0, 0), "@@program_name@@", bg="0A2A52")
value_cell(tbl1.cell(1, 0), "Service Fee", "@@down_payment@@ TL + VAT")
value_cell(tbl1.cell(2, 0), "Success Bonus", "%@@success_bonus@@ + VAT", bg="FFFFFF")

add_bonus_note(doc1)
add_notes_footer(doc1)
doc1.save(os.path.join(OUT, "Pricing_1Program_template.docx"))
print("OK Pricing_1Program_template.docx")


# ─── TEMPLATE 2 — Two programs ────────────────────────────────────
doc2 = make_doc()
add_title(doc2)
add_client(doc2)

tbl2 = doc2.add_table(rows=3, cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Table Grid"
colors2 = ["0A2A52", "3A6B35"]
for ci, (nm, fee, bon) in enumerate([
    ("@@program_name@@",  "@@down_payment@@",  "@@success_bonus@@"),
    ("@@program2_name@@", "@@program2_fee@@",  "@@program2_bonus@@"),
]):
    header_cell(tbl2.cell(0, ci), nm, bg=colors2[ci])
    value_cell(tbl2.cell(1, ci),  "Service Fee",    f"{fee} TL + VAT")
    value_cell(tbl2.cell(2, ci),  "Success Bonus",  f"%{bon} + VAT", bg="FFFFFF")

add_bonus_note(doc2)
add_notes_footer(doc2)
doc2.save(os.path.join(OUT, "Pricing_2Programs_template.docx"))
print("OK Pricing_2Programs_template.docx")


# ─── TEMPLATE 3 — Three programs ──────────────────────────────────
doc3 = make_doc()
add_title(doc3)
add_client(doc3)

tbl3 = doc3.add_table(rows=3, cols=3)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = "Table Grid"
colors3 = ["0A2A52", "3A6B35", "1A5F9E"]
for ci, (nm, fee, bon) in enumerate([
    ("@@program_name@@",  "@@down_payment@@",  "@@success_bonus@@"),
    ("@@program2_name@@", "@@program2_fee@@",  "@@program2_bonus@@"),
    ("@@program3_name@@", "@@program3_fee@@",  "@@program3_bonus@@"),
]):
    header_cell(tbl3.cell(0, ci), nm, bg=colors3[ci])
    value_cell(tbl3.cell(1, ci),  "Service Fee",   f"{fee} TL + VAT")
    value_cell(tbl3.cell(2, ci),  "Success Bonus", f"%{bon} + VAT", bg="FFFFFF")

add_bonus_note(doc3)
add_notes_footer(doc3)
doc3.save(os.path.join(OUT, "Pricing_3Programs_template.docx"))
print("OK Pricing_3Programs_template.docx")

print("\nDone! Upload each file via Contracts > Manage Templates.")
